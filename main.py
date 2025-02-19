
from docx import Document
from app import *
from db import *





#Spring 2025 OBHR 33000-008 LEC
def main(): 
    
    document = doc_init()
    app = QApplication()
# Create a Qt widget, which will be our window.
    window = MainWindow()
    window.show()  # IMPORTANT!!!!! Windows are hidden by default.
    
    # Start the event loop.
    app.exec()
    file = open("temp.txt", "a")
    for db in app.py.tables:
        file.write(db + "\n")
        
    file.close()
    # document.save(name + ".docx")

def doc_init():
    
    name = " "
    print(name)
    document = Document()
    document.add_heading
    document.add_heading(name, 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading(name, level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    
    return document


if __name__ == "__main__":
    main()