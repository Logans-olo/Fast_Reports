from docx import Document

# Function to create and save a new document
def doc_init(name):
    document = Document()
    document.add_heading(name + "'s Report", 0)  # Add a heading at level 0 (largest)
    document.save(name + '.docx')  # Save the document with the given name

# Function to add a paragraph to the document
def doc_add_paragraph(docName, text):
    document = Document(docName)  # Open the existing document
    document.add_paragraph(text)  # Add the paragraph with the given text
    print("saved")
    document.save(docName)  # Save the updated document

# Function to add a table to the document
def doc_add_table(docName, colNames, data):
    document = Document(docName)  # Open the existing document
    table = document.add_table(rows=1, cols=len(colNames))  # Create table with headers
    # Add the header row
    hdr_cells = table.rows[0].cells
    for i, colName in enumerate(colNames):
        hdr_cells[i].text = colName

    # Add data rows
    for row_data in data:
        row = table.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = str(value)
    
    document.save(docName)  # Save the updated document

# Function to add a heading to the document
def doc_add_heading(docName, text):
    document = Document(docName)  # Open the existing document
    document.add_heading(text, 2)  # Add a heading with the specified level (1 = largest)
    document.save(docName)  # Save the updated document
